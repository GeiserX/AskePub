# core_worker.py - AskePub Core Worker
# Generic ePub study assistant: parse, annotate, and export study notes.

import os
import copy
import hashlib
import logging
import sqlite3
import subprocess
import pytz
from datetime import datetime

import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from openai import OpenAI

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
TIMEZONE = pytz.timezone("Europe/Madrid")
CACHE_DB_PATH = "/app/dbs/cache.db"
OPENAI_MODEL = "gpt-4o-mini"

# Minimum plain-text length for a chapter to be considered real content
# (filters out cover pages, copyright notices, etc.)
MIN_CHAPTER_TEXT_LENGTH = 200


# ===================================================================
# 1. parse_epub
# ===================================================================

def parse_epub(epub_path: str) -> list[dict]:
    """Read an ePub file and return a list of chapter dicts.

    Each dict has:
        id        - the item id/href inside the ePub
        title     - best-effort chapter title
        content_html - raw HTML of the chapter body
        content_text - plain text extracted from the HTML
    """
    logger.info("parse_epub - path: %s", epub_path)
    book = epub.read_epub(epub_path, options={"ignore_ncx": False})

    # Build a map from href -> toc title using the TOC
    toc_title_map: dict[str, str] = {}
    _walk_toc(book.toc, toc_title_map)

    chapters: list[dict] = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        raw_html = item.get_content().decode("utf-8", errors="replace")
        soup = BeautifulSoup(raw_html, "html.parser")
        plain_text = soup.get_text(separator="\n", strip=True)

        # Skip very short items (covers, copyright, etc.)
        if len(plain_text) < MIN_CHAPTER_TEXT_LENGTH:
            continue

        # Determine title: prefer TOC entry, fall back to first <h1>-<h3>, then item id
        href_key = item.get_name()
        title = toc_title_map.get(href_key, "")
        if not title:
            heading = soup.find(["h1", "h2", "h3"])
            title = heading.get_text(strip=True) if heading else href_key

        chapters.append({
            "id": href_key,
            "title": title,
            "content_html": raw_html,
            "content_text": plain_text,
        })

    logger.info("parse_epub - found %d chapters", len(chapters))
    return chapters


def _walk_toc(toc_items, title_map: dict):
    """Recursively walk the ePub TOC and populate href->title mapping."""
    for item in toc_items:
        if isinstance(item, tuple):
            # Nested section: (Section, [children])
            section, children = item
            if hasattr(section, "href") and hasattr(section, "title"):
                clean_href = section.href.split("#")[0]
                title_map[clean_href] = section.title
            _walk_toc(children, title_map)
        elif hasattr(item, "href") and hasattr(item, "title"):
            clean_href = item.href.split("#")[0]
            title_map[clean_href] = item.title


# ===================================================================
# 2. query_openai
# ===================================================================

def _init_cache_db():
    """Ensure the cache SQLite database and table exist."""
    os.makedirs(os.path.dirname(CACHE_DB_PATH), exist_ok=True)
    conn = sqlite3.connect(CACHE_DB_PATH)
    conn.execute(
        "CREATE TABLE IF NOT EXISTS cache "
        "(hash TEXT PRIMARY KEY, response TEXT)"
    )
    conn.commit()
    conn.close()


def _cache_key(chapter_text: str, questions: list[str]) -> str:
    payload = chapter_text + "|||" + "|||".join(questions)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def query_openai(
    book_title: str,
    chapter_title: str,
    chapter_text: str,
    questions: list[str],
    language: str,
) -> dict[int, str]:
    """Send chapter text + questions to OpenAI and return numbered answers.

    Returns a dict mapping question index (0-based) to the answer string.
    Responses are cached in SQLite to avoid redundant API calls.
    """
    logger.info("query_openai - book: %s, chapter: %s", book_title, chapter_title)

    # --- Check cache ---
    _init_cache_db()
    key = _cache_key(chapter_text, questions)
    conn = sqlite3.connect(CACHE_DB_PATH)
    row = conn.execute("SELECT response FROM cache WHERE hash = ?", (key,)).fetchone()
    if row:
        logger.info("query_openai - cache hit for chapter: %s", chapter_title)
        conn.close()
        return _parse_numbered_answers(row[0], len(questions))
    conn.close()

    # --- Build prompt ---
    numbered_questions = "\n".join(
        f"{i + 1}. {q}" for i, q in enumerate(questions)
    )

    system_prompt = (
        f"You are a study assistant helping prepare notes for a book. "
        f"The book is titled '{book_title}'. "
        f"For the chapter '{chapter_title}', answer each of the following study questions "
        f"based on the chapter content. Number your answers to match the questions. "
        f"Be concise but insightful. "
        f"Respond entirely in the language with code '{language}'."
    )

    user_message = (
        f"Chapter content:\n\n{chapter_text}\n\n"
        f"Questions:\n{numbered_questions}"
    )

    # --- Call OpenAI ---
    client = OpenAI()
    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        temperature=0.4,
    )

    answer_text = response.choices[0].message.content.strip()
    logger.info("query_openai - received response for chapter: %s", chapter_title)

    # --- Store in cache ---
    conn = sqlite3.connect(CACHE_DB_PATH)
    conn.execute(
        "INSERT OR REPLACE INTO cache (hash, response) VALUES (?, ?)",
        (key, answer_text),
    )
    conn.commit()
    conn.close()

    return _parse_numbered_answers(answer_text, len(questions))


def _parse_numbered_answers(text: str, n_questions: int) -> dict[int, str]:
    """Parse a numbered response into a dict {0: answer0, 1: answer1, ...}.

    Tries to split on lines starting with "1.", "2.", etc.  Falls back to
    splitting by double newlines if numbering is absent.
    """
    import re

    answers: dict[int, str] = {}
    # Try splitting by numbered pattern "1." "2." etc.
    parts = re.split(r"\n(?=\d+\.\s)", text)
    for part in parts:
        m = re.match(r"(\d+)\.\s*(.*)", part, re.DOTALL)
        if m:
            idx = int(m.group(1)) - 1  # 0-based
            answers[idx] = m.group(2).strip()

    if len(answers) >= n_questions:
        return answers

    # Fallback: split by double newline
    chunks = [c.strip() for c in text.split("\n\n") if c.strip()]
    for i, chunk in enumerate(chunks[:n_questions]):
        if i not in answers:
            # Remove leading number if present
            cleaned = re.sub(r"^\d+\.\s*", "", chunk)
            answers[i] = cleaned

    return answers


# ===================================================================
# 3. write_annotated_epub
# ===================================================================

def write_annotated_epub(
    original_epub_path: str,
    chapters_with_notes: list[dict],
    output_path: str,
) -> str:
    """Create a new ePub with study notes appended to selected chapters.

    chapters_with_notes: list of dicts with at least
        {"id": str, "questions": [str], "notes": {idx: str}}

    Returns the output_path.
    """
    logger.info("write_annotated_epub - source: %s, output: %s", original_epub_path, output_path)
    book = epub.read_epub(original_epub_path, options={"ignore_ncx": False})

    # Map chapter id -> notes payload
    notes_map: dict[str, dict] = {
        ch["id"]: ch for ch in chapters_with_notes if ch.get("notes")
    }

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        href = item.get_name()
        if href not in notes_map:
            continue

        ch = notes_map[href]
        questions = ch.get("questions", [])
        notes = ch.get("notes", {})
        if not notes:
            continue

        # Build the notes HTML block
        notes_html = _build_notes_html(questions, notes)

        # Append to existing content
        raw_html = item.get_content().decode("utf-8", errors="replace")
        # Insert before closing </body> if present, otherwise just append
        if "</body>" in raw_html:
            raw_html = raw_html.replace("</body>", notes_html + "\n</body>")
        else:
            raw_html += notes_html

        item.set_content(raw_html.encode("utf-8"))

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    epub.write_epub(output_path, book)
    logger.info("write_annotated_epub - done: %s", output_path)
    return output_path


def _build_notes_html(questions: list[str], notes: dict[int, str]) -> str:
    """Return an HTML snippet with the study notes section."""
    lines = [
        '<div style="margin-top: 2em; padding: 1em; '
        'background-color: #FFF8DC; border-top: 2px solid #CCC;">',
        '<h3 style="margin-top: 0;">Study Notes - AskePub</h3>',
    ]
    for idx in sorted(notes.keys()):
        q_text = questions[idx] if idx < len(questions) else f"Question {idx + 1}"
        a_text = notes[idx]
        lines.append(
            f'<p><strong>{_escape_html(q_text)}</strong></p>'
            f'<p>{_escape_html(a_text)}</p>'
        )
    lines.append("</div>")
    return "\n".join(lines)


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# ===================================================================
# 4. write_docx_pdf
# ===================================================================

def write_docx_pdf(
    book_title: str,
    chapters_data: list[dict],
    telegram_user: str,
) -> tuple[str, str]:
    """Create a DOCX (and PDF via abiword) with all study notes.

    chapters_data: list of {"title": str, "questions": [str], "notes": {idx: str}}

    Returns (docx_path, pdf_path).
    """
    logger.info("write_docx_pdf - book: %s, user: %s", book_title, telegram_user)

    now_date = datetime.now(TIMEZONE).strftime("%Y-%m-%d")
    user_dir = f"userBackups/{telegram_user}"
    os.makedirs(user_dir, exist_ok=True)

    safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in book_title)[:60]
    docx_path = os.path.join(user_dir, f"askepub-{safe_title}-{now_date}.docx")
    pdf_path = os.path.join(user_dir, f"askepub-{safe_title}-{now_date}.pdf")

    document = Document()

    # Add bold list number style
    bold_style = document.styles.add_style("Bold List Number", WD_STYLE_TYPE.PARAGRAPH)
    bold_style.font.bold = True

    # Title
    document.add_heading(book_title, level=0)

    # Credit line
    p = document.add_paragraph()
    p.add_run("Generated by AskePub - https://github.com/GeiserX/AskePub").italic = True

    # Chapters
    for ch in chapters_data:
        document.add_heading(ch["title"], level=1)
        questions = ch.get("questions", [])
        notes = ch.get("notes", {})

        for idx in sorted(notes.keys()):
            q_text = questions[idx] if idx < len(questions) else f"Question {idx + 1}"
            a_text = notes[idx]

            q_para = document.add_paragraph(style="Bold List Number")
            q_para.add_run(q_text).font.size = Pt(12)
            document.add_paragraph(a_text)

    document.save(docx_path)
    logger.info("write_docx_pdf - DOCX saved: %s", docx_path)

    # Convert to PDF
    cmd = f"xvfb-run abiword --to=pdf --to-name='{pdf_path}' '{docx_path}'"
    try:
        subprocess.run(cmd, shell=True, check=True, timeout=120)
        logger.info("write_docx_pdf - PDF saved: %s", pdf_path)
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
        logger.error("write_docx_pdf - PDF conversion failed: %s", exc)

    return docx_path, pdf_path


# ===================================================================
# 5. main (orchestrator)
# ===================================================================

def main(
    epub_path: str,
    telegram_user: str,
    selected_chapter_ids: list[str],
    questions: list[str],
    language: str,
) -> tuple[str, str, str]:
    """Orchestrate the full workflow.

    1. Parse the ePub
    2. Filter to selected chapters
    3. Query OpenAI for each chapter
    4. Write annotated ePub
    5. Write DOCX + PDF
    6. Return (annotated_epub_path, docx_path, pdf_path)
    """
    logger.info(
        "main - epub: %s, user: %s, chapters: %s, language: %s",
        epub_path, telegram_user, selected_chapter_ids, language,
    )

    # 1. Parse
    all_chapters = parse_epub(epub_path)

    # 2. Filter
    selected = [ch for ch in all_chapters if ch["id"] in selected_chapter_ids]
    if not selected:
        logger.warning("main - no matching chapters found, processing all")
        selected = all_chapters

    # Derive book title from ePub metadata
    book = epub.read_epub(epub_path, options={"ignore_ncx": False})
    book_title = "Unknown"
    dc_title = book.get_metadata("DC", "title")
    if dc_title:
        book_title = dc_title[0][0]
    logger.info("main - book title: %s", book_title)

    # 3. Query OpenAI for each selected chapter
    chapters_with_notes: list[dict] = []
    for ch in selected:
        notes = query_openai(
            book_title=book_title,
            chapter_title=ch["title"],
            chapter_text=ch["content_text"],
            questions=questions,
            language=language,
        )
        chapters_with_notes.append({
            "id": ch["id"],
            "title": ch["title"],
            "questions": questions,
            "notes": notes,
        })

    # 4. Write annotated ePub
    now_date = datetime.now(TIMEZONE).strftime("%Y-%m-%d")
    user_dir = f"userBackups/{telegram_user}"
    os.makedirs(user_dir, exist_ok=True)

    safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in book_title)[:60]
    annotated_epub_path = os.path.join(user_dir, f"askepub-{safe_title}-{now_date}.epub")
    write_annotated_epub(epub_path, chapters_with_notes, annotated_epub_path)

    # 5. Write DOCX + PDF
    docx_path, pdf_path = write_docx_pdf(book_title, chapters_with_notes, telegram_user)

    # 6. Return paths
    logger.info("main - complete. Files: %s, %s, %s", annotated_epub_path, docx_path, pdf_path)
    return annotated_epub_path, docx_path, pdf_path


if __name__ == "__main__":
    # Example usage
    result = main(
        epub_path="example.epub",
        telegram_user="test_user",
        selected_chapter_ids=[],
        questions=[
            "What is the main argument of this chapter?",
            "Provide a concrete example or illustration for a key point.",
            "How does this chapter connect to the overall theme of the book?",
        ],
        language="en",
    )
    print("Output files:", result)
